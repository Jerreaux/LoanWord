from django.shortcuts import render
from django.http import HttpResponse, FileResponse
from .models import Document
from docx import Document as DocxDocument
import os
from django.conf import settings
import uuid

# Create your views here.

def upload_document(request):
    if request.method == 'POST' and request.FILES.get('document'):
        # Ambil file yang diupload
        uploaded_file = request.FILES['document']
        
        # Buat instance Document dan simpan
        document = Document(file=uploaded_file)
        document.save()
        
        # Dapatkan path file yang diupload
        input_path = document.file.path
        
        # Buat nama file output yang unik
        output_filename = f"Jerreaux_{uuid.uuid4().hex[:8]}.docx"
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        
        # Pastikan direktori processed ada
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Proses dokumen
        doc = DocxDocument(input_path)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.italic = True
        
        # Simpan dokumen yang telah dimodifikasi
        doc.save(output_path)
        
        # Simpan path file yang telah diproses ke dalam model
        document.processed_file = f'processed/{output_filename}'
        document.save()
        
        # Redirect ke halaman download
        return render(request, 'download.html', {'document_id': document.id})
    
    return render(request, 'upload.html')

def download_document(request, document_id):
    try:
        document = Document.objects.get(id=document_id)
        file_path = document.processed_file.path
        
        if os.path.exists(file_path):
            # Buka file dan buat FileResponse
            response = FileResponse(open(file_path, 'rb'), as_attachment=True)
            # Set nama file untuk download
            response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
            return response
        else:
            return HttpResponse("File tidak ditemukan", status=404)
    except Document.DoesNotExist:
        return HttpResponse("Dokumen tidak ditemukan", status=404)
    except Exception as e:
        return HttpResponse(f"Terjadi kesalahan: {str(e)}", status=500)

